# tests/test_extractor_docx.py
from pathlib import Path
from docx import Document
from src.adapters.extractors.docx import extract_docx


def test_extract_docx(tmp_path):
    p = tmp_path / "x.docx"
    d = Document()
    d.add_paragraph("Первый абзац")
    d.add_paragraph("Второй абзац")
    d.save(p)
    text = extract_docx(p)
    assert "Первый абзац" in text
    assert "Второй абзац" in text
